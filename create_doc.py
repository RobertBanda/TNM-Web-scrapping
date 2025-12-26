from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('TNM Broadband Data Analysis Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Comprehensive Guide to TNM-Web-scrapping Project')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0].font
subtitle_format.italic = True
subtitle_format.size = Pt(12)

doc.add_paragraph()  # Blank line

# Overview Section
doc.add_heading('Overview', 1)
overview_para = doc.add_paragraph(
    'This project performs comprehensive analysis of TNM (Telekom Networks Malawi) broadband customer data, '
    'following an 8-step analytical framework. The analysis provides valuable business insights and actionable '
    'recommendations for improving customer retention and revenue growth.'
)

doc.add_paragraph('The analysis includes:', style='List Bullet')
doc.add_paragraph('Wikipedia Data Scraping: Scrapes TNM information from Wikipedia', style='List Bullet 2')
doc.add_paragraph('Data Generation: Creates a realistic TNM broadband customer dataset', style='List Bullet 2')
doc.add_paragraph('Data Analysis: Performs Exploratory Data Analysis (EDA) following the 8-step structure with interactive visualizations', style='List Bullet 2')
doc.add_paragraph('Business Insights: Provides actionable recommendations for TNM', style='List Bullet 2')

doc.add_paragraph()
note_para = doc.add_paragraph('Note: ', style='Normal')
note_run = note_para.add_run('The dataset contains sample data for analysis. The original dataset contained over 10 million customer records.')
note_run.bold = True

doc.add_paragraph()

# Files Section
doc.add_heading('Project Files', 1)
doc.add_paragraph('The project consists of the following key files:', style='List Bullet')
doc.add_paragraph('scrape_tnm_wikipedia.py - Scrapes TNM information from Wikipedia', style='List Bullet 2')
doc.add_paragraph('generate_tnm_data.py - Generates TNM broadband customer dataset', style='List Bullet 2')
doc.add_paragraph('tnm_analysis.py - Main analysis script (8-step framework)', style='List Bullet 2')
doc.add_paragraph('tnm_predictions.py - Predictive analysis: What-if scenarios & 5-year forecast', style='List Bullet 2')
doc.add_paragraph('run_tnm_analysis.py - Master script to run the entire pipeline', style='List Bullet 2')
doc.add_paragraph('requirements.txt - Python dependencies', style='List Bullet 2')
doc.add_paragraph('tnm_broadband.csv - Generated customer dataset (1000 rows sample)', style='List Bullet 2')

doc.add_paragraph()

# Installation Section
doc.add_heading('Installation', 1)
doc.add_paragraph('1. Install required packages:', style='Normal')
code_para = doc.add_paragraph('pip install -r requirements.txt', style='Normal')
code_run = code_para.runs[0]
code_run.font.name = 'Consolas'
code_run.font.size = Pt(10)

doc.add_paragraph()

# Usage Section
doc.add_heading('Usage', 1)

doc.add_heading('Option 1: Run Complete Pipeline', 2)
code_para2 = doc.add_paragraph('python run_tnm_analysis.py', style='Normal')
code_run2 = code_para2.runs[0]
code_run2.font.name = 'Consolas'
code_run2.font.size = Pt(10)

doc.add_paragraph('This will:', style='Normal')
doc.add_paragraph('1. Scrape Wikipedia data (optional)', style='List Number')
doc.add_paragraph('2. Generate the TNM dataset', style='List Number')
doc.add_paragraph('3. Run the complete analysis', style='List Number')

doc.add_heading('Option 2: Run Individual Steps', 2)

doc.add_heading('1. Generate Dataset', 3)
code_para3 = doc.add_paragraph('python generate_tnm_data.py', style='Normal')
code_run3 = code_para3.runs[0]
code_run3.font.name = 'Consolas'
code_run3.font.size = Pt(10)
doc.add_paragraph('Creates a customer dataset with configurable number of records (default: 1000 for the sample).', style='Normal')

doc.add_heading('2. Run Analysis', 3)
code_para4 = doc.add_paragraph('python tnm_analysis.py', style='Normal')
code_run4 = code_para4.runs[0]
code_run4.font.name = 'Consolas'
code_run4.font.size = Pt(10)
doc.add_paragraph('Performs the 8-step exploratory data analysis.', style='Normal')

doc.add_heading('3. Run Predictive Analysis', 3)
code_para5 = doc.add_paragraph('python tnm_predictions.py', style='Normal')
code_run5 = code_para5.runs[0]
code_run5.font.name = 'Consolas'
code_run5.font.size = Pt(10)
doc.add_paragraph('This script performs predictive analysis including:', style='Normal')
doc.add_paragraph('• Predicts outcomes if TNM fixes network downtime issues', style='List Bullet')
doc.add_paragraph('• Predicts impact of customer plan upgrades', style='List Bullet')
doc.add_paragraph('• Forecasts business metrics for the next 5 years', style='List Bullet')
doc.add_paragraph('• Compares scenarios with and without improvements', style='List Bullet')

doc.add_paragraph()

# Analysis Structure Section
doc.add_heading('Analysis Structure', 1)
doc.add_paragraph('The analysis follows a structured 8-step framework:', style='Normal')

doc.add_heading('Step 1: Import Libraries & Data', 2)
doc.add_paragraph('Loads the TNM broadband dataset into a pandas DataFrame for analysis.', style='Normal')

doc.add_heading('Step 2: Understand the Data', 2)
doc.add_paragraph('Provides dataset information, statistics, and business understanding framework to familiarize with the data structure.', style='Normal')

doc.add_heading('Step 3: Data Quality Checks', 2)
doc.add_paragraph('Performs data quality validation including:', style='Normal')
doc.add_paragraph('• Missing values check', style='List Bullet')
doc.add_paragraph('• Duplicate detection', style='List Bullet')

doc.add_heading('Step 4: Data Preparation', 2)
doc.add_paragraph('Prepares the data for analysis by:', style='Normal')
doc.add_paragraph('• Converting date columns to proper datetime format', style='List Bullet')
doc.add_paragraph('• Creating derived columns like monthly_revenue', style='List Bullet')

doc.add_heading('Step 5: Exploratory Data Analysis (EDA)', 2)
doc.add_paragraph('Performs comprehensive exploratory analysis including:', style='Normal')
doc.add_paragraph('• Revenue by Subscription Plan analysis', style='List Bullet')
doc.add_paragraph('• Data Usage vs Churn correlation analysis', style='List Bullet')
doc.add_paragraph('• Network Downtime vs Churn impact analysis', style='List Bullet')

doc.add_heading('Step 6: Key Business Questions', 2)
doc.add_paragraph('Answers critical business questions such as:', style='Normal')
doc.add_paragraph('• Which region experiences the highest churn?', style='List Bullet')
doc.add_paragraph('• Which subscription plan is most stable (lowest churn)?', style='List Bullet')

doc.add_heading('Step 7: Business Insights', 2)
doc.add_paragraph('Provides key insights including:', style='Normal')
doc.add_paragraph('• Network downtime impact on customer churn', style='List Bullet')
doc.add_paragraph('• Plan performance analysis', style='List Bullet')
doc.add_paragraph('• Usage pattern insights', style='List Bullet')
doc.add_paragraph('• Regional analysis and recommendations', style='List Bullet')

doc.add_heading('Step 8: Recommendations', 2)
doc.add_paragraph('Provides actionable recommendations for:', style='Normal')
doc.add_paragraph('• Network stability improvements', style='List Bullet')
doc.add_paragraph('• Plan upgrade strategies', style='List Bullet')
doc.add_paragraph('• KPI monitoring protocols', style='List Bullet')
doc.add_paragraph('• Customer retention initiatives', style='List Bullet')

doc.add_paragraph()

# Dataset Schema Section
doc.add_heading('Dataset Schema', 1)
doc.add_paragraph('The TNM broadband dataset includes the following columns:', style='Normal')

schema_items = [
    ('customer_id', 'Unique customer identifier'),
    ('region', 'Customer region (Northern, Central, Southern, Lilongwe, Blantyre)'),
    ('subscription_plan', 'Plan type (Basic, Standard, Premium, Unlimited)'),
    ('signup_date', 'Customer signup date'),
    ('monthly_fee', 'Monthly subscription fee in MWK (Malawian Kwacha)'),
    ('data_usage_gb', 'Monthly data usage in GB'),
    ('network_downtime_hours', 'Network downtime hours per month'),
    ('churned', 'Churn status (Yes/No)')
]

for field, desc in schema_items:
    para = doc.add_paragraph(f'{field}: ', style='Normal')
    run = para.add_run(desc)
    run.font.italic = True

doc.add_paragraph()

# Key Findings Section
doc.add_heading('Key Findings', 1)

findings = [
    ('Network Downtime', 'Higher downtime directly correlates with increased customer churn. Customers experiencing frequent network outages are significantly more likely to discontinue service.'),
    ('Plan Performance', 'Unlimited plans generate the highest revenue and demonstrate the lowest churn rates, indicating strong customer satisfaction and value perception.'),
    ('Usage Patterns', 'Low-usage customers are at higher risk of churning, suggesting that engagement and usage are important factors in customer retention.'),
    ('Regional Issues', 'The Southern region requires urgent network infrastructure improvements due to higher downtime and churn rates compared to other regions.')
]

for title, desc in findings:
    doc.add_heading(title, 2)
    doc.add_paragraph(desc, style='Normal')

doc.add_paragraph()

# Recommendations Section
doc.add_heading('Recommendations', 1)

recs = [
    'Improve network stability in Southern region through infrastructure upgrades',
    'Target Basic & Standard plan users with upgrade offers to Premium/Unlimited plans',
    'Monitor downtime KPIs closely with real-time dashboards and alert systems',
    'Introduce loyalty incentives for low-usage customers to increase engagement'
]

for rec in recs:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_paragraph()

# PREDICTIVE ANALYSIS SECTION - DETAILED EXPLANATION
doc.add_heading('Predictive Analysis: Detailed Explanation', 1)
doc.add_paragraph(
    'The predictive analysis module (tnm_predictions.py) is one of the most powerful components of this project. '
    'It uses statistical modeling and scenario analysis to predict future business outcomes based on different improvement strategies.',
    style='Normal'
)

doc.add_heading('How Predictive Analysis Works', 2)

doc.add_heading('1. Churn Probability Calculation', 3)
doc.add_paragraph(
    'The system uses a mathematical model to calculate the probability that a customer will churn (cancel their service). '
    'This probability is based on multiple factors that influence customer satisfaction and retention.',
    style='Normal'
)

doc.add_paragraph('The churn probability formula considers:', style='Normal')
doc.add_paragraph('• Base churn probability: A starting point of 10% (0.1) representing inherent churn risk', style='List Bullet')
doc.add_paragraph('• Network downtime impact: Higher downtime increases churn probability (up to 40% additional risk)', style='List Bullet')
doc.add_paragraph('• Data usage influence: Higher usage decreases churn probability (up to 20% risk reduction)', style='List Bullet')
doc.add_paragraph('• Regional factors: Southern region has an additional 15% churn risk penalty due to infrastructure issues', style='List Bullet')

formula_para = doc.add_paragraph('Mathematical Formula:', style='Normal')
formula_run = formula_para.runs[0]
formula_run.bold = True
formula_run.font.size = Pt(12)

code_para6 = doc.add_paragraph(
    'Churn_Probability = Base_Probability + (Downtime/50) - (Usage/200) + Regional_Penalty\n'
    'Final probability is clamped between 5% and 70% to ensure realistic values',
    style='Normal'
)
code_run6 = code_para6.runs[0]
code_run6.font.name = 'Consolas'
code_run6.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('2. What-If Scenarios Analysis', 3)
doc.add_paragraph(
    'The predictive analysis runs four different improvement scenarios to understand the potential impact of various strategies. '
    'Each scenario simulates changes to the business and calculates the resulting customer retention and revenue impact.',
    style='Normal'
)

doc.add_heading('Scenario 1: Fix Network Downtime in Southern Region', 4)
doc.add_paragraph(
    'This scenario simulates reducing network downtime in the Southern region by 50%, bringing it closer to the levels experienced '
    'in other regions. The analysis then recalculates churn probabilities for all customers in the Southern region with the improved '
    'downtime values.',
    style='Normal'
)
doc.add_paragraph('How it works:', style='Normal')
doc.add_paragraph('1. Identifies all customers in the Southern region', style='List Number')
doc.add_paragraph('2. Reduces their network_downtime_hours by 50%', style='List Number')
doc.add_paragraph('3. Recalculates churn probability for each affected customer using the new downtime values', style='List Number')
doc.add_paragraph('4. Calculates the difference in expected churn rate', style='List Number')
doc.add_paragraph('5. Estimates revenue impact by multiplying retained customers by average monthly fee', style='List Number')

doc.add_paragraph()

doc.add_heading('Scenario 2: Fix All Network Downtime', 4)
doc.add_paragraph(
    'This scenario simulates a comprehensive network improvement program that reduces downtime across all regions by 60%. '
    'This represents a more ambitious infrastructure investment that would improve service quality nationwide.',
    style='Normal'
)
doc.add_paragraph('How it works:', style='Normal')
doc.add_paragraph('1. Reduces network_downtime_hours by 60% for all customers across all regions', style='List Number')
doc.add_paragraph('2. Recalculates churn probabilities for all customers', style='List Number')
doc.add_paragraph('3. Estimates total customer retention and revenue impact', style='List Number')
doc.add_paragraph('4. Provides a baseline for understanding the maximum potential impact of network improvements', style='List Number')

doc.add_paragraph()

doc.add_heading('Scenario 3: Customer Plan Upgrades', 4)
doc.add_paragraph(
    'This scenario simulates an upgrade campaign that moves customers from Basic and Standard plans to higher-tier plans (Standard and Premium respectively). '
    'Research shows that customers on higher-tier plans typically have lower churn rates and generate more revenue.',
    style='Normal'
)
doc.add_paragraph('How it works:', style='Normal')
doc.add_paragraph('1. Identifies customers on Basic and Standard plans', style='List Number')
doc.add_paragraph('2. Simulates upgrading them to Standard and Premium plans respectively', style='List Number')
doc.add_paragraph('3. Increases their monthly fees to the new plan rates', style='List Number')
doc.add_paragraph('4. Increases their data usage (customers on better plans typically use more data)', style='List Number')
doc.add_paragraph('5. Recalculates churn probabilities (higher usage reduces churn risk)', style='List Number')
doc.add_paragraph('6. Calculates revenue increase from both plan upgrades and reduced churn', style='List Number')

doc.add_paragraph()
revenue_para = doc.add_paragraph('Revenue Impact Calculation:', style='Normal')
revenue_run = revenue_para.runs[0]
revenue_run.bold = True

doc.add_paragraph(
    'Revenue Increase = (Upgrade Fee Increase × Number of Upgraded Customers) + (Retained Customers × Average Monthly Fee)',
    style='Normal'
)

doc.add_paragraph()

doc.add_heading('Scenario 4: Combined Improvements', 4)
doc.add_paragraph(
    'This is the most comprehensive scenario that combines all improvement strategies. It simulates fixing network downtime '
    'AND upgrading customer plans simultaneously. This represents the optimal business strategy if TNM implements all recommended improvements.',
    style='Normal'
)
doc.add_paragraph('How it works:', style='Normal')
doc.add_paragraph('1. Applies Scenario 1 (Southern region downtime reduction)', style='List Number')
doc.add_paragraph('2. Applies Scenario 2 (All regions downtime reduction)', style='List Number')
doc.add_paragraph('3. Applies Scenario 3 (Customer plan upgrades)', style='List Number')
doc.add_paragraph('4. Calculates the combined impact on churn and revenue', style='List Number')
doc.add_paragraph('5. Provides the maximum potential business impact', style='List Number')

doc.add_paragraph()

doc.add_heading('3. 5-Year Business Forecast', 3)
doc.add_paragraph(
    'The forecasting module projects TNM\'s business metrics (customers and revenue) over the next 5 years under two scenarios: '
    'maintaining the current state versus implementing the recommended improvements.',
    style='Normal'
)

doc.add_heading('Forecasting Methodology', 4)
doc.add_paragraph('The forecast uses a month-by-month projection model that accounts for:', style='Normal')
doc.add_paragraph('• Customer growth: Assumes 2% monthly customer growth rate (conservative estimate)', style='List Bullet')
doc.add_paragraph('• Customer churn: Applies the calculated churn rate each month', style='List Bullet')
doc.add_paragraph('• Revenue calculation: Based on number of active customers × average monthly fee', style='List Bullet')

doc.add_paragraph()
formula_para2 = doc.add_paragraph('Forecasting Formula (Month by Month):', style='Normal')
formula_run2 = formula_para2.runs[0]
formula_run2.bold = True

code_para7 = doc.add_paragraph(
    'New_Customers = Current_Customers × Growth_Rate (2%)\n'
    'Churned_Customers = Current_Customers × Churn_Rate\n'
    'Next_Month_Customers = Current_Customers + New_Customers - Churned_Customers\n'
    'Monthly_Revenue = Active_Customers × Average_Monthly_Fee',
    style='Normal'
)
code_run7 = code_para7.runs[0]
code_run7.font.name = 'Consolas'
code_run7.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Scenario A: No Improvements (Current State)', 4)
doc.add_paragraph(
    'This forecast maintains the current churn rate and business conditions. It shows what would happen if TNM continues '
    'operating without making any improvements to network infrastructure or customer retention strategies.',
    style='Normal'
)
doc.add_paragraph('Key assumptions:', style='Normal')
doc.add_paragraph('• Current churn rate continues unchanged', style='List Bullet')
doc.add_paragraph('• No network improvements implemented', style='List Bullet')
doc.add_paragraph('• No plan upgrade campaigns', style='List Bullet')
doc.add_paragraph('• Natural market growth continues at 2% monthly', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Scenario B: With Improvements', 4)
doc.add_paragraph(
    'This forecast uses the results from Scenario 4 (Combined Improvements) as the starting point. It shows the business impact '
    'if TNM implements all recommended improvements immediately and maintains the improved churn rate going forward.',
    style='Normal'
)
doc.add_paragraph('Key assumptions:', style='Normal')
doc.add_paragraph('• Reduced churn rate from combined improvements', style='List Bullet')
doc.add_paragraph('• Network improvements implemented and maintained', style='List Bullet')
doc.add_paragraph('• Plan upgrades completed', style='List Bullet')
doc.add_paragraph('• Higher average revenue per user (ARPU) due to plan upgrades', style='List Bullet')
doc.add_paragraph('• Same 2% monthly growth rate, but with better retention', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Forecast Comparison and Business Impact', 4)
doc.add_paragraph(
    'By comparing the two scenarios, the analysis reveals the cumulative business impact of implementing improvements. '
    'The difference between Scenario A and Scenario B shows:', style='Normal'
)
doc.add_paragraph('• Additional customers retained over 5 years', style='List Bullet')
doc.add_paragraph('• Additional revenue generated over 5 years', style='List Bullet')
doc.add_paragraph('• Return on investment (ROI) for improvement initiatives', style='List Bullet')
doc.add_paragraph('• Strategic value of proactive improvements', style='List Bullet')

doc.add_paragraph()

doc.add_heading('4. Output Files and Visualizations', 3)
doc.add_paragraph('The predictive analysis generates several output files:', style='Normal')

outputs = [
    ('tnm_scenario_comparison.png', 'Visual comparison chart showing churn rates and revenue increases for all scenarios'),
    ('tnm_5year_forecast.png', 'Line charts showing customer and revenue projections over 5 years for both scenarios'),
    ('detailed_predictions.json', 'Complete data in JSON format with all scenario results and forecast details'),
    ('scenario_comparison.csv', 'Tabular data comparing all scenarios in spreadsheet format'),
    ('5year_forecast.csv', 'Year-by-year forecast data for both scenarios'),
    ('predictions_summary.txt', 'Human-readable text summary of all predictions and recommendations')
]

for filename, description in outputs:
    para = doc.add_paragraph(f'{filename}: ', style='Normal')
    run = para.add_run(description)
    run.font.italic = True

doc.add_paragraph()

doc.add_heading('Understanding the Results', 3)
doc.add_paragraph(
    'When reviewing the prediction results, it\'s important to understand what the numbers mean:',
    style='Normal'
)

doc.add_heading('Churn Rate Reduction', 4)
doc.add_paragraph(
    'The churn rate reduction shows how many percentage points the churn rate would decrease. For example, if the current '
    'churn rate is 12% and a scenario shows a reduction of 2 percentage points, the new churn rate would be 10%. This means '
    '2 out of every 100 customers who would have churned will now stay with TNM.',
    style='Normal'
)

doc.add_heading('Customers Retained', 4)
doc.add_paragraph(
    'This shows the estimated number of customers who would not churn if the improvements are implemented. These are customers '
    'who would have left under current conditions but will stay with improved service quality.',
    style='Normal'
)

doc.add_heading('Revenue Increase', 4)
doc.add_paragraph(
    'The revenue increase has two components:',
    style='Normal'
)
doc.add_paragraph('1. Revenue from retained customers: These customers continue paying their monthly fees', style='List Number')
doc.add_paragraph('2. Revenue from plan upgrades: Customers who upgrade pay higher monthly fees', style='List Number')
doc.add_paragraph(
    'The total revenue increase shows the additional monthly and annual revenue TNM would generate by implementing improvements.',
    style='Normal'
)

doc.add_paragraph()

doc.add_heading('5. Statistical Sampling and Performance', 3)
doc.add_paragraph(
    'For efficiency with large datasets, the prediction analysis uses statistical sampling. When the dataset contains more than '
    '50,000 records, it samples 50,000 records for scenario calculations, then scales the results to the full dataset size. '
    'This approach maintains accuracy while significantly improving processing speed.',
    style='Normal'
)

doc.add_paragraph()
note_para2 = doc.add_paragraph('Important Note: ', style='Normal')
note_run2 = note_para2.add_run(
    'The sampled predictions are statistically scaled to represent the full dataset, ensuring that the results accurately '
    'reflect the impact on all customers, not just the sample.'
)
note_run2.bold = True

doc.add_paragraph()

doc.add_heading('6. Practical Applications', 3)
doc.add_paragraph('The predictive analysis results can be used for:', style='Normal')
doc.add_paragraph('• Business planning and budgeting for improvement initiatives', style='List Bullet')
doc.add_paragraph('• ROI calculations for network infrastructure investments', style='List Bullet')
doc.add_paragraph('• Strategic decision-making about which improvements to prioritize', style='List Bullet')
doc.add_paragraph('• Stakeholder presentations demonstrating the value of improvements', style='List Bullet')
doc.add_paragraph('• Setting targets and KPIs for improvement programs', style='List Bullet')
doc.add_paragraph('• Long-term business forecasting and financial planning', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Conclusion: Value of Predictive Analysis', 2)
doc.add_paragraph(
    'The predictive analysis module transforms raw data into actionable business intelligence. By quantifying the impact of '
    'different improvement strategies, it enables data-driven decision-making. The what-if scenarios help TNM understand the '
    'potential returns on investment before committing resources, while the 5-year forecast provides a strategic roadmap for '
    'sustainable business growth.',
    style='Normal'
)

doc.add_paragraph()
doc.add_paragraph(
    'By implementing the recommended improvements, TNM can expect significant improvements in customer retention, revenue growth, '
    'and overall business performance over the next 5 years.',
    style='Normal'
)

doc.add_paragraph()
doc.add_page_break()

# Output Files Section
doc.add_heading('Output Files', 1)
doc.add_paragraph('After running the analysis, you\'ll get the following output files:', style='Normal')

output_list = [
    ('tnm_broadband.csv', 'Customer dataset (1000 rows sample)'),
    ('tnm_revenue_by_plan.png', 'Revenue visualization chart'),
    ('tnm_data_usage_vs_churn.png', 'Usage vs churn analysis chart'),
    ('tnm_downtime_vs_churn.png', 'Downtime impact visualization'),
    ('tnm_churn_by_region.png', 'Regional churn analysis chart'),
    ('tnm_churn_by_plan.png', 'Plan stability analysis chart'),
    ('tnm_scenario_comparison.png', 'Comparison of improvement scenarios'),
    ('tnm_5year_forecast.png', '5-year business forecast charts'),
    ('tnm_wikipedia_data.json', 'Scraped Wikipedia data (if available)')
]

for filename, description in output_list:
    para = doc.add_paragraph(f'{filename}: ', style='Normal')
    run = para.add_run(description)
    run.font.italic = True

doc.add_paragraph()

# Requirements Section
doc.add_heading('Technical Requirements', 1)
doc.add_paragraph('The project requires the following:', style='Normal')

reqs = [
    'Python 3.7 or higher',
    'pandas - for data manipulation and analysis',
    'matplotlib - for creating visualizations',
    'seaborn - for enhanced statistical visualizations',
    'numpy - for numerical computations',
    'beautifulsoup4 - for web scraping',
    'requests - for HTTP requests',
    'python-docx - for generating documentation (optional)'
]

for req in reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph()

# Notes Section
doc.add_heading('Important Notes', 1)
notes = [
    'The dataset sample contains 1,000 customer records for demonstration purposes',
    'The original full dataset was designed to contain 10+ million records',
    'The dataset is generated based on realistic TNM operations in Malawi',
    'MWK (Malawian Kwacha) is used as the currency throughout the analysis',
    'Analysis follows business analytics best practices',
    'All visualizations are displayed interactively and saved as high-resolution PNG files',
    'Data generation uses vectorized operations for optimal performance',
    'Predictive analysis uses statistical sampling for efficiency with large datasets'
]

for note in notes:
    doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()

# Footer
doc.add_paragraph()
footer_para = doc.add_paragraph('TNM-Web-scrapping Project Documentation')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.runs[0]
footer_run.font.size = Pt(10)
footer_run.font.italic = True

# Save the document
output_file = 'TNM_Documentation.docx'
doc.save(output_file)
print(f'Documentation created successfully: {output_file}')

